import openpyxl
from docx import Document

def extract_text_from_excel(excel_file, docx_file):
    wb = openpyxl.load_workbook(excel_file)
    sheet = wb.active

    document = Document()

    for row in sheet.iter_rows(values_only=True):
        doc_type, title, category, question, content = row[:5]

        document.add_page_break()

        p = document.add_paragraph(f"[{doc_type}]->[{title}]->[{category}]->{question}\n\n")
        p.add_run(content).bold = True

    document.save(docx_file)

excel_file = "sample.xlsx"
docx_file = "output.docx"

extract_text_from_excel(excel_file, docx_file)
